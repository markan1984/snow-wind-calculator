from flask import send_file
from docx import Document
from docx.shared import Inches


def snow1_to_docx(slope, Ce, Ct, city, mui, level, snow_area, Sk, Sk_formula, s):
    #Creation of the docx file
    document = Document()
    document.add_heading('Определение снеговых нагрузок на односкатные покрытия.', level=1)
    document.add_paragraph('Сбор снеговой нагрузки ведем согласно указаниям СН 2.01.04.')
    document.add_paragraph('Согласно указаниям национального приложения НП.1 на территории РБ применимы только нормальные условия, установленные в соответствии с 3.2(1).')
    document.add_paragraph('Характеристическое значение снеговой нагрузки на покрытие для постоянных/переходных расчетных ситуаций следует определять согласно п.5.2(3):')
    p1 = document.add_paragraph("")
    p1.add_run('s=μ₁∙cₑ∙cₜ∙sₖ').bold = True
    document.add_paragraph('Где:')
    document.add_paragraph(f'μ₁={mui} - коэффициент формы снеговых нагрузок (п.5.3) при α={slope}° и с учетом п.п.5.3.2(2);', style='List')
    document.add_picture('./static/snow_pictures/snow_load1.jpg', width=Inches(2))
    document.add_paragraph(f'cₑ={Ce} - коэффициент окружающей среды (табл. НП.1.3(BY));', style='List')
    document.add_paragraph(f'cₜ={Ct} - температурный коэффициент (п.5.2(8) НП.1(BY));', style='List')
    document.add_paragraph(f'sₖ={Sk_formula}{float(Sk/1000)}кПа - характеристическое значение снеговых нагрузок на грунт для района г.{city} (район {snow_area}) при абсолютной отметке местности А={level};', style='List')
    document.add_paragraph('Тогда:')
    p2 = document.add_paragraph("")
    p2.add_run(f's=μ₁∙cₑ∙cₜ∙sₖ={mui}∙{Ce}∙{Ct}∙{float(Sk/1000)}={s}кПа').bold = True
    document.add_paragraph(f'Итого: характеристическое значение нагрузки на покрытие кровли составляет s={s}кПа')
    #Save .doc file:
    document.save('./calculations/snow1.docx')
    #Download created .doc file:
    return send_file('./calculations/snow1.docx',as_attachment=True)


def snow2_to_docx(slope1, slope2, Ce, Ct, city, mui1, mui2, level, snow_area, Sk, Sk_formula, s1, s2):
    #Creation of the docx file
    document = Document()
    document.add_heading('Определение снеговых нагрузок на двускатные покрытия.', level=1)
    document.add_paragraph('Сбор снеговой нагрузки ведем согласно указаниям СН 2.01.04.')
    document.add_paragraph('Согласно указаниям национального приложения НП.1 на территории РБ применимы только нормальные условия, установленные в соответствии с 3.2(1).')
    document.add_paragraph('Характеристическое значение снеговой нагрузки на покрытие для постоянных/переходных расчетных ситуаций следует определять согласно п.5.2(3):')
    p1 = document.add_paragraph("")
    p1.add_run('s=μᵢ∙cₑ∙cₜ∙sₖ').bold = True
    document.add_paragraph('Где:')
    document.add_paragraph(f'μ₁(α₁)={mui1} - коэффициент формы снеговых нагрузок для первого ската кровли при α₁={slope1}° и с учетом п.5.3.3(2);', style='List')
    document.add_paragraph(f'μ₁(α₂)={mui2} - коэффициент формы снеговых нагрузок для второго ската кровли при α₂={slope2}° и с учетом п.5.3.3(2);', style='List')
    document.add_picture('./static/snow_pictures/snow_load2.jpg', width=Inches(3.5))
    document.add_paragraph(f'cₑ={Ce} - коэффициент окружающей среды (табл. НП.1.3(BY));', style='List')
    document.add_paragraph(f'cₜ={Ct} - температурный коэффициент (п.5.2(8) НП.1(BY));', style='List')
    document.add_paragraph(f'sₖ={Sk_formula}{float(Sk/1000)}кПа - характеристическое значение снеговых нагрузок на грунт для района г.{city} (район {snow_area}) при абсолютной отметке местности А={level};', style='List')
    document.add_paragraph('Тогда:')
    p2 = document.add_paragraph("")
    p2.add_run(f's₁=μ₁∙cₑ∙cₜ∙sₖ={mui1}∙{Ce}∙{Ct}∙{float(Sk/1000)}={s1}кПа').bold = True
    p2.add_run('- для первого ската').bold = False
    p3 = document.add_paragraph("")
    p3.add_run(f's₂=μ₁∙cₑ∙cₜ∙sₖ={mui2}∙{Ce}∙{Ct}∙{float(Sk/1000)}={s2}кПа').bold = True
    p3.add_run('- для второго ската').bold = False
    document.add_paragraph(f'Итого: характеристическое значение снеговой нагрузки на первый и второй скаты кровли соответственно будут равны: s₁={s1}кПа, s₂={s2}кПа.')
    #Save .doc file:
    document.save('./calculations/snow2.docx')
    #Download created .doc file:
    return send_file('./calculations/snow2.docx',as_attachment=True)


def snow3_to_docx(slope1, slope2, Ce, Ct, city, mui1, mui1a, mui2, mui2a, mu2, level, snow_area, Sk, Sk_formula, s1, s1a, s2, s2a, s3):
    document = Document()
    document.add_heading('Определение снеговых нагрузок на многопролетные покрытия.', level=1)
    document.add_paragraph('Сбор снеговой нагрузки ведем согласно указаниям СН 2.01.04.')
    document.add_paragraph('Согласно указаниям национального приложения НП.1 на территории РБ применимы только нормальные условия, установленные в соответствии с 3.2(1).')
    document.add_paragraph('Характеристическое значение снеговой нагрузки на покрытие для постоянных/переходных расчетных ситуаций следует определять согласно п.5.2(3):')
    p1 = document.add_paragraph("")
    p1.add_run('s=μ₁∙cₑ∙cₜ∙sₖ').bold = True
    document.add_paragraph('Где:')
    document.add_paragraph(f'μ₁(α₁)={mui1} - коэффициент формы снеговых нагрузок для первого ската кровли при α₁={slope1}° и с учетом п.5.3.4(2);', style='List')
    document.add_paragraph(f'μ₁ₐ(α₂)={mui2a} - коэффициент формы снеговых нагрузок для второго (внутреннего) ската кровли при α₂={slope2}° и с учетом примыкающего ската;', style='List')
    document.add_paragraph(f'μ₁ₐ(α₁)={mui1a} - коэффициент формы снеговых нагрузок для первого (внутреннего) ската кровли при α₁={slope1}° и с учетом примыкающего ската;', style='List')
    document.add_paragraph(f'μ₁(α₂)={mui2} - коэффициент формы снеговых нагрузок для второго ската кровли при α₂={slope2}° и с учетом п.5.3.4(2);', style='List')
    document.add_paragraph(f'μ₂(α)={mu2} - коэффициент формы снеговых нагрузок для центральной части кровли с учетом заноса снега при α=(α₁+α₂)/2 и с учетом таблицы 5.2 и п. 5.3.4(3);', style='List')
    document.add_paragraph(f'cₑ={Ce} - коэффициент окружающей среды (табл. НП.1.3(BY));', style='List')
    document.add_paragraph(f'cₜ={Ct} - температурный коэффициент (п.5.2(8) НП.1(BY));', style='List')
    document.add_paragraph(f'sₖ={Sk_formula}{float(Sk/1000)}кПа - характеристическое значение снеговых нагрузок на грунт для района г.{city} (район {snow_area}) при абсолютной отметке местности А={level};', style='List')
    document.add_paragraph('В расчете следует рассматривать 2 случая распределения снеговых нагрузок: без учета заноса снега и с учетом заноса.')
    document.add_picture('./static/snow_pictures/snow_load3.jpg', width=Inches(5))
    ##document.add_page_break()
    p2 = document.add_paragraph("Для")
    p2.add_run(' случая I ').bold = True
    p2.add_run('(без учета заноса снега) схема распределения снеговых нагрузок будет иметь вид:').bold = False
    document.add_picture('./static/snow_pictures/snow_load3-1.jpg', width=Inches(3.5))
    document.add_paragraph('Где:', style='List')

    p3 = document.add_paragraph("")
    p3.add_run(f's₁=μ₁∙cₑ∙cₜ∙sₖ={mui1}∙{Ce}∙{Ct}∙{float(Sk/1000)}={s1}кПа').bold = True
    p3.add_run(f'– на скат 1(с α₁={slope1}°)').bold = False

    p9 = document.add_paragraph("")
    p9.add_run(f's₂ₐ=μ₁ₐ∙cₑ∙cₜ∙sₖ={mui2a}∙{Ce}∙{Ct}∙{float(Sk/1000)}={s2a}кПа').bold = True
    p9.add_run(f'– на скат 2(с α₂={slope2}°)').bold = False

    p10 = document.add_paragraph("")
    p10.add_run(f's₁ₐ=μ₁ₐ∙cₑ∙cₜ∙sₖ={mui1a}∙{Ce}∙{Ct}∙{float(Sk/1000)}={s1a}кПа').bold = True
    p10.add_run(f'– на скат 1(с α₁={slope1}°)').bold = False

    p4 = document.add_paragraph("")
    p4.add_run(f's₂=μ₁∙cₑ∙cₜ∙sₖ={mui2}∙{Ce}∙{Ct}∙{float(Sk/1000)}={s2}кПа').bold = True
    p4.add_run(f'– на скат 2(с α₂={slope2}°)').bold = False

    p5 = document.add_paragraph("Для")
    p5.add_run(' случая II ').bold = True
    p5.add_run('(с учётом заноса снега) схема распределения снеговых нагрузок будет иметь вид:').bold = False
    document.add_picture('./static/snow_pictures/snow_load3-2.jpg', width=Inches(3.5))
    document.add_paragraph('Где:', style='List')
    p6 = document.add_paragraph("")
    p6.add_run(f's₁=μ₁∙cₑ∙cₜ∙sₖ={mui1}∙{Ce}∙{Ct}∙{float(Sk/1000)}={s1}кПа').bold = True
    p6.add_run(f'– на крайний скат 1(с α₁={slope1})').bold = False
    p8 = document.add_paragraph("")
    p8.add_run(f's₃=μ₂∙cₑ∙cₜ∙sₖ={mu2}∙{Ce}∙{Ct}∙{float(Sk/1000)}={s3}кПа').bold = True
    p8.add_run('– на центральную часть кровли (с α=(α₁+α₂)/2)').bold = False
    p7 = document.add_paragraph("")
    p7.add_run(f's₂=μ₁∙cₑ∙cₜ∙sₖ={mui2}∙{Ce}∙{Ct}∙{float(Sk/1000)}={s2}кПа').bold = True
    p7.add_run(f'– на крайний скат 2(с α₂={slope2}°)').bold = False

    #Save .doc file:
    document.save('./calculations/snow3.docx')
    #Download created .doc file:
    return send_file('./calculations/snow3.docx',as_attachment=True)


def snow4_to_docx(b, h, Ce, Ct, city, mu3, level, snow_area, Sk, Sk_formula, s1, s2):
    document = Document()
    document.add_heading('Определение снеговых нагрузок на арочные покрытия.', level=1)
    document.add_paragraph('Сбор снеговой нагрузки ведем согласно указаниям СН 2.01.04.')
    document.add_paragraph('Согласно указаниям национального приложения НП.1 на территории РБ применимы только нормальные условия, установленные в соответствии с 3.2(1).')
    document.add_paragraph('Характеристическое значение снеговой нагрузки на покрытие для постоянных/переходных расчетных ситуаций следует определять согласно п.5.2(3):')
    p1 = document.add_paragraph("")
    p1.add_run('s=μ₁∙cₑ∙cₜ∙sₖ').bold = True
    document.add_paragraph('Где:')
    document.add_paragraph(f'μ₁(I)=0,8 - коэффициент формы снеговых нагрузок для арочного покрытия без учета заносов снега п.5.3.5(2);', style='List')
    document.add_paragraph(f'μ₃(II)=0,2+10h/b=0,2+10*{h}/{b} [но не более 2,0] = {mu3} - коэффициент формы снеговых нагрузок для арочного покрытия с учетом заносов снега согласно п.5.3.5(3);', style='List')
    document.add_paragraph(f'cₑ={Ce} - коэффициент окружающей среды (табл. НП.1.3(BY));', style='List')
    document.add_paragraph(f'cₜ={Ct} - температурный коэффициент (п.5.2(8) НП.1(BY));', style='List')
    document.add_paragraph(f'sₖ={Sk_formula}{float(Sk/1000)}кПа - характеристическое значение снеговых нагрузок на грунт для района г.{city} (район {snow_area}) при абсолютной отметке местности А={level};', style='List')
    document.add_paragraph('В расчете следует рассматривать 2 случая распределения снеговых нагрузок: без учета заноса снега и с учетом заноса.')
    document.add_picture('./static/snow_pictures/snow_load4.jpg', width=Inches(3))
    document.add_page_break()
    p2 = document.add_paragraph("Для")
    p2.add_run(' случая I ').bold = True
    p2.add_run('(без учета заноса снега) схема распределения снеговых нагрузок будет иметь вид:').bold = False
    document.add_picture('./static/snow_pictures/snow_load4-1.jpg', width=Inches(2.5))
    document.add_paragraph('Где:', style='List')
    p3 = document.add_paragraph("")
    p3.add_run(f's=μ₁∙cₑ∙cₜ∙sₖ=0,8∙{Ce}∙{Ct}∙{float(Sk/1000)}={s1}кПа').bold = True
    p3.add_run(f'- на всю длину зоны ls (при β≤60°)').bold = False

    p5 = document.add_paragraph("Для")
    p5.add_run(' случая II ').bold = True
    p5.add_run('(с учётом заноса снега) схема распределения снеговых нагрузок будет иметь вид:').bold = False
    document.add_picture('./static/snow_pictures/snow_load4-2.jpg', width=Inches(2.5))
    document.add_paragraph('Где:', style='List')
    p6 = document.add_paragraph("")
    p6.add_run(f's=μ₃∙cₑ∙cₜ∙sₖ={mu3}∙{Ce}∙{Ct}∙{float(Sk/1000)}={s2}кПа').bold = True

    #Save .doc file:
    document.save('./calculations/snow4.docx')
    #Download created .doc file:
    return send_file('./calculations/snow4.docx',as_attachment=True)


def snow5_to_docx(b1,b2,h,L,S,slope1, slope2, Ce, Ct, city, mui1, mui2, mus, muw, mu2, ls, p_link1, p_link2, p_link3, level, snow_area, Sk, Sk_formula, s1, s2, s3, sx):
    b1 = round(float(b1/1000),2)
    b2 = round(float(b2/1000),2)
    h = round(float(h/1000),2)
    L = round(float(L/1000),2)
    mu2 = round(mu2,3)
    document = Document()
    document.add_heading('Определение снеговых нагрузок на покрытие кровли на перепаде высот.', level=1)
    document.add_paragraph('Сбор снеговой нагрузки ведем согласно указаниям СН 2.01.04.')
    document.add_paragraph('Согласно указаниям национального приложения НП.1 на территории РБ применимы только нормальные условия, установленные в соответствии с 3.2(1).')
    document.add_paragraph('В расчете следует рассматривать два случая распределения снеговых нагрузок: без учета заноса снега и с учетом заноса.')
    document.add_picture(f'.{p_link1}', width=Inches(3))
    p1 = document.add_paragraph("Где:")
    p1.add_run(f'b₁={b1}м, b₂={b2}м, h={h}м, α₁={slope1}°, α₂={slope2}°, ls={ls}м.').bold = True
    document.add_paragraph('Коэффициенты формы снеговой нагрузки, которые необходимо применять для покрытий, примыкающих к более высоким сооружениям рассчитывают по следующим формулам:')

    p6 = document.add_paragraph("",style='List')
    p6.add_run(f'μ₁(α₂)={mui2}').bold = True
    p6.add_run(' - коэффициент формы снеговых нагрузок для нижележащего покрытия без учета заносов снега но с учетом таблицы 5.2 и п.5.3.2(2);').bold = False

    p7 = document.add_paragraph("",style='List')
    p7.add_run(f'μ₂=μs+μw={mus}+{muw}={mu2}').bold = True
    p7.add_run(' - коэффициент формы снеговых нагрузок для нижележащего покрытия с учетом заносов и сползания снега;').bold = False

    document.add_paragraph('Где:',style='List')

    p8 = document.add_paragraph("",style='List')
    p8.add_run(f'μs={mus}').bold = True
    p8.add_run(' - коэффициент формы снеговой нагрузки, учитывающий соскальзывание снега с вышележащей кровли согласно п. 5.3.6(1);').bold = False

    p9 = document.add_paragraph("",style='List')
    p9.add_run('μw').bold = True
    p9.add_run(' - коэффициент формы снеговой нагрузки, учитывающий влияние ветра:').bold = False

    p10 = document.add_paragraph("",style='List')
    p10.add_run(f'μw=(b1+b2)/2h<γh/sₖ ➔ ➔  ({b1}+{b2})/(2∙{h}) < 2∙{h}/{float(Sk/1000)}.').bold = True

    p11 = document.add_paragraph("",style='List')
    p11.add_run(f'μw={muw}').bold = True
    p11.add_run(' - учитывая площадь нижележащей кровли ').bold = False
    p11.add_run(f'S=b2xL={b2}x{L}={S}м²').bold = True
    p11.add_run(' и требования п.5.3.6(1) НП.1;').bold = False

    p12 = document.add_paragraph("",style='List')
    p12.add_run(f'sₖ={Sk_formula}{float(Sk/1000)}кПа').bold = True
    p12.add_run(f' - характеристическое значение снеговых нагрузок на грунт для района г.{city} (район {snow_area}) при абсолютной отметке местности А={level};').bold = False

    p13 = document.add_paragraph("",style='List')
    p13.add_run(f'cₑ={Ce}').bold = True
    p13.add_run(f' - коэффициент окружающей среды (табл. НП.1.3(BY));').bold = False

    p14 = document.add_paragraph("",style='List')
    p14.add_run(f'cₜ={Ct}').bold = True
    p14.add_run(' - температурный коэффициент (п.5.2(8) НП.1(BY));').bold = False

    p2 = document.add_paragraph("")
    p2.add_run('Для случая I').bold = True
    p2.add_run('(без учета заноса снега) схема распределения снеговых нагрузок будет иметь вид:').bold = False
    document.add_picture(f'.{p_link2}', width=Inches(2.5))
    document.add_paragraph('Где:',style='List')
    p3 = document.add_paragraph("")
    p3.add_run(f's=μ₁∙cₑ∙cₜ∙sₖ={mui2}∙{Ce}∙{Ct}∙{float(Sk/1000)}={s2}кПа').bold = True
    p3.add_run('(с учетом уклона пониженной кровли α₂=3° и с учетом п.5.3.2(2)).').bold = False

    p4 = document.add_paragraph("")
    p4.add_run('Для случая II').bold = True
    p4.add_run('(c учетом заноса снега) схема распределения снеговых нагрузок будет иметь вид:').bold = False
    document.add_picture(f'.{p_link3}', width=Inches(2.5))
    document.add_paragraph('Где:',style='List')

    p5 = document.add_paragraph("",style='List')
    p5.add_run(f'ls={ls}м ').bold = True
    p5.add_run('- длина зоны повышенных снегоотложений').bold = False

    p15 = document.add_paragraph("",style='List')
    p15.add_run(f's₃=μ₂∙cₑ∙cₜ∙sₖ = {mu2}∙{Ce}∙{Ct}∙{float(Sk/1000)}={s3}кПа.').bold = True

    p16 = document.add_paragraph("",style='List')
    p16.add_run(f's₂={sx}кПа.').bold = True

    p17 = document.add_paragraph("",style='List')
    p17.add_run(f's₁=μ₁∙cₑ∙cₜ∙sₖ = {mui2}∙{Ce}∙{Ct}∙{float(Sk/1000)}={s2}кПа.').bold = True

    #Save .doc file:
    document.save('./calculations/snow5.docx')
    #Download created .doc file:
    return send_file('./calculations/snow5.docx',as_attachment=True)


def snow6_to_docx(h, ls, Ce, Ct, mu1, mu2, Sk, s1, s2):
    #Creation of the docx file
    document = Document()
    document.add_heading('Определение снеговых нагрузок у надстроек и заграждений кровли', level=1)
    document.add_paragraph('Сбор снеговой нагрузки ведем согласно указаниям СН 2.01.04.')
    document.add_paragraph('Cхема распределения снеговых нагрузок будет иметь вид:')
    document.add_picture('./static/snow_pictures/snow_load6.jpg', width=Inches(2))
    document.add_paragraph('Где:')

    p1 = document.add_paragraph("")
    p1.add_run(f'h={h}м').bold = True
    p1.add_run(f' - высота надстройки или заграждения;').bold = False

    p2 = document.add_paragraph("")
    p2.add_run(f'ls={ls}м').bold = True
    p2.add_run(f' - длина зоны повышенных снегоотложений;').bold = False

    p3 = document.add_paragraph("")
    p3.add_run(f's₁=μ₁∙cₑ∙cₜ∙sₖ = {mu1}∙{Ce}∙{Ct}∙{Sk}={s1}кПа').bold = True
    #p3.add_run(f' (согласно п.6.2(2))').bold = False

    p4 = document.add_paragraph("")
    p4.add_run(f's₂=μ₂∙cₑ∙cₜ∙sₖ = {mu2}∙{Ce}∙{Ct}∙{Sk}={s2}кПа').bold = True
    p4.add_run(f' (согласно п.6.2(2))').bold = False

    #Save .doc file:
    document.save('./calculations/snow6.docx')
    #Download created .doc file:
    return send_file('./calculations/snow6.docx',as_attachment=True)


def snow7_to_docx(slope, b, sinalfa, s, Fs):
    #Creation of the docx file
    document = Document()
    document.add_heading('Определение снеговых нагрузок на снегоудерживающие заграждения.', level=1)
    document.add_paragraph('Сбор снеговой нагрузки ведем согласно указаниям СН 2.01.04.')
    document.add_paragraph('Cхема снеговых нагрузок на снегоудерживающие заграждения будет иметь вид:')
    document.add_picture('./static/snow_pictures/snow_slope7.jpg', width=Inches(2))
    document.add_paragraph('Где:')

    p1 = document.add_paragraph("")
    p1.add_run(f'α={slope}°;').bold = True
    p1.add_run(f' - уклон кровли;').bold = False

    p2 = document.add_paragraph("")
    p2.add_run(f'b={b}м').bold = True
    p2.add_run(f' - горизонтальное расстояние от снегоудерживающей решетки или надстройки до следующей решетки или до конька;').bold = False

    p3 = document.add_paragraph("")
    p3.add_run(f'Fs=s∙b∙sin({slope}°)={s}∙{b}∙{sinalfa}={Fs}кН/м').bold = True
    p3.add_run(f' - Искомая касательная нагрузка от снега (в плоскости скольжения) на 1м.п. заграждения.').bold = False

    #Save .doc file:
    document.save('./calculations/snow7.docx')
    #Download created .doc file:
    return send_file('./calculations/snow7.docx',as_attachment=True)